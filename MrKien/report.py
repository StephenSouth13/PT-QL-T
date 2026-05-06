import pandas as pd
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_report():
    # 1. Khởi tạo Document
    doc = Document()
    
    # --- STYLE ---
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)

    # --- TIÊU ĐỀ ---
    title = doc.add_heading('ĐỀ ÁN CÁ NHÂN: PHÂN TÍCH TÀI CHÍNH ĐỊNH LƯỢNG', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # --- THÔNG TIN SINH VIÊN ---
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run('Họ và tên: Mạnh Hồ Kiên\nMSSV: 33241026046\nGiảng viên hướng dẫn: Lê Văn Lâm')
    run.font.size = Pt(14)
    run.bold = True

    doc.add_page_break()

    # --- CÂU 1 ---
    doc.add_heading('CÂU 1: PHÂN TÍCH KỸ THUẬT (CASE STUDY: TCB)', level=1)
    doc.add_paragraph(
        "Sử dụng bộ dữ liệu OHLCV của Techcombank (TCB) trong giai đoạn 2024-2026, "
        "nghiên cứu thực hiện giải mã các tín hiệu xung lực thông qua hệ thống chỉ báo đa lớp."
    )
    
    # Chèn ảnh PTKT
    try:
        doc.add_picture('Q1_PTKT_TCB.png', width=Inches(6))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    except:
        doc.add_paragraph("[Lỗi: Không tìm thấy ảnh Q1_PTKT_TCB.png]")

    doc.add_paragraph(
        "Dựa trên biểu đồ, TCB đang vận hành trong một cấu trúc Uptrend bền vững. "
        "Đường giá duy trì vị thế phía trên SMA(20), xác nhận xu hướng tăng trung hạn. "
        "Chỉ báo MACD và RSI cho thấy dòng tiền thông minh đang thẩm thấu tốt vào cổ phiếu, "
        "duy trì đà tăng trưởng mà chưa xuất hiện dấu hiệu quá mua cực đoan."
    )

    # --- CÂU 2 ---
    doc.add_heading('CÂU 2: QUẢN TRỊ DANH MỤC VÀ KIỂM ĐỊNH CAPM', level=1)
    
    # (A) Thống kê
    doc.add_heading('(A) Đặc tính thống kê và Tính chuẩn', level=2)
    doc.add_paragraph("Bảng thống kê mô tả cho 5 mã cổ phiếu (TCB, MWG, REE, GAS, HPG):")
    
    # Đọc dữ liệu từ Excel để đưa vào Word (nếu muốn làm bảng xịn)
    try:
        stats_df = pd.read_excel("Bao_cao_MrKien.xlsx", sheet_name='Stats')
        table = doc.add_table(rows=1, cols=len(stats_df.columns))
        hdr_cells = table.rows[0].cells
        for i, col in enumerate(stats_df.columns):
            hdr_cells[i].text = str(col)
        
        for index, row in stats_df.iterrows():
            row_cells = table.add_row().cells
            for i, val in enumerate(row):
                row_cells[i].text = f"{val:.4f}" if isinstance(val, float) else str(val)
        table.style = 'Table Grid'
    except:
        doc.add_paragraph("[Lỗi: Không đọc được dữ liệu Stats từ Excel]")

    doc.add_picture('Q2A_Histogram.png', width=Inches(4.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph(
        "Nhận xét: Hệ số Kurtosis dương (Excess Kurtosis > 0) khẳng định hiện tượng đuôi dày (Fat-tails). "
        "Điều này bác bỏ giả thuyết phân phối chuẩn Gaussian, hàm ý thị trường tiềm ẩn rủi ro cực đoan cao."
    )

    # (B) Markowitz
    doc.add_heading('(B) Lý thuyết danh mục Markowitz', level=2)
    doc.add_paragraph(
        "Thông qua ma trận hiệp phương sai, nghiên cứu xác lập tỷ trọng MVP cho cặp TCB-MWG "
        "nhằm tối ưu hóa biên độ an toàn cho nhà đầu tư. Việc đa dạng hóa sang 5 ngành giúp "
        "giảm thiểu rủi ro phi hệ thống một cách hiệu quả."
    )

    # (C) CAPM
    doc.add_heading('(C) Mô hình CAPM và Đường SML', level=2)
    doc.add_picture('Q2C_SML.png', width=Inches(5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph(
        "Kiểm định CAPM cho thấy Beta của TCB phản ánh đúng tính chất của một cổ phiếu Aggressive. "
        "Kết quả hồi quy OLS xác nhận hệ số Alpha không có ý nghĩa thống kê (p > 0.05), "
        "cho thấy cổ phiếu được định giá sát với giá trị lý thuyết của mô hình."
    )

    # --- PHỤ LỤC R CODE ---
    doc.add_page_break()
    doc.add_heading('PHỤ LỤC: MÃ NGUỒN R (Dành cho kiểm định)', level=1)
    rcode = (
        "library(quantmod)\n"
        "symbols <- c('TCB.VN', 'MWG.VN', 'REE.VN', 'GAS.VN', 'HPG.VN')\n"
        "getSymbols(symbols, src='yahoo', from='2024-04-10', to='2026-04-10')\n"
        "R <- na.omit(diff(log(Cl(merge(TCB.VN, MWG.VN, REE.VN, GAS.VN, HPG.VN)))))\n"
        "chartSeries(TCB.VN); addMACD(); addRSI(); addBBands()\n"
        "fit <- lm((R[,1]-0.03/250) ~ (rowMeans(R)-0.03/250))\n"
        "summary(fit)"
    )
    code_para = doc.add_paragraph()
    run_code = code_para.add_run(rcode)
    run_code.font.name = 'Courier New'
    run_code.font.size = Pt(10)

    # LƯU FILE
    doc.save("BAO_CAO_CUOI_KY_MANH_HO_KIEN.docx")
    print("--- BÁO CÁO ĐÃ TẠO THÀNH CÔNG! ---")

if __name__ == "__main__":
    create_report()